# ''' From page 307 in Automate the Boring Stuff
#     Reading Word Documents, initial examples  '''


# import docx
# doc = docx.Document('demo.docx') # make sure this existing doc is in the same cd you are in
# x = len(doc.paragraphs) # to see how many 'Paragraph' objects there are in the doc
# print(f'There are {x} paragraph objects in the document.')

# # print(doc.paragraphs[0].text) # Print the very first 'Paragraph' object in the document
# # print(doc.paragraphs[1].text) # Print the second

# len(doc.paragraphs[1].runs) # Print the first 'Run' within the first 'Paragraph object' - aka, a run is like a chunk of text, and italics and bold will start a new 'Run'


'''Getting the full text from a .docx file
- With this program, you can now import it in another file (see doc.py)
(see page 308 in book if you need more clairification)
'''
import docx

def getText(filename): # getText() will open the Word Doct
    doc = docx.Document(filename)
    fullText = []
    for para in doc.paragraphs: # loop over all the Paragraph objects, and then append it to the empty list in previous line. 
        fullText.append(para.text)
    return '\n'.join(fullText) # Join them together with a new line, so it's not all one giant chunk of text
