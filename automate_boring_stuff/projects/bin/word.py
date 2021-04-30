''' From page 307 in Automate the Boring Stuff
    Reading Word Documents, initial examples  '''


import docx
doc = docx.Document('demo.docx') # make sure this existing doc is in the same cd you are in
x = len(doc.paragraphs) # to see how many 'Paragraph' objects there are in the doc
print(f'There are {x} paragraph objects in the document.')

# print(doc.paragraphs[0].text) # Print the very first 'Paragraph' object in the document
# print(doc.paragraphs[1].text) # Print the second

len(doc.paragraphs[1].runs) # Print the first 'Run' within the first 'Paragraph object' - aka, a run is like a chunk of text, and italics and bold will start a new 'Run'